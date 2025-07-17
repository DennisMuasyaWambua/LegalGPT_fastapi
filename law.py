"""
Website Vectorizer: Crawl, process, vectorize website content and downloadable files
Optimized for Kenya Law site (https://new.kenyalaw.org)
"""
import os
import requests
from bs4 import BeautifulSoup
import urllib.parse
import mimetypes
import tempfile
import asyncio
from concurrent.futures import ThreadPoolExecutor
import logging
from typing import List, Dict, Any, Set, Tuple, Optional
import time
import re
import random
import json
from pathlib import Path
from dataclasses import dataclass, asdict
from urllib.robotparser import RobotFileParser

# For vectorization
import numpy as np
from sentence_transformers import SentenceTransformer

# For vector database
import chromadb
from chromadb.config import Settings

# Text extraction libraries
import PyPDF2
from docx import Document
import pandas as pd
import openpyxl
import csv

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# File handler to save logs to disk
file_handler = logging.FileHandler("crawler.log")
file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
logger.addHandler(file_handler)

@dataclass
class CrawlerStats:
    """Statistics for the crawler"""
    pages_visited: int = 0
    files_downloaded: int = 0
    failed_pages: int = 0
    failed_downloads: int = 0
    vectorized_texts: int = 0
    start_time: float = 0.0
    
    def print_stats(self):
        """Print current statistics"""
        elapsed = time.time() - self.start_time if self.start_time > 0 else 0
        logger.info(f"Stats: {self.pages_visited} pages, {self.files_downloaded} files, "
                   f"{self.failed_pages} failed pages, {self.failed_downloads} failed downloads, "
                   f"{self.vectorized_texts} texts vectorized, "
                   f"Runtime: {elapsed:.1f}s")

class WebsiteVectorizer:
    def __init__(
        self, 
        base_url: str,
        model_name: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        max_pages: int = 500,
        max_depth: int = 5,
        vector_db_path: str = "./vector_db",
        download_dir: str = None,
        concurrent_requests: int = 8,
        request_delay: float = 0.5,
        user_agent: str = None,
        respect_robots_txt: bool = True,
        proxy: str = None,
        persist_downloads: bool = False,
        collection = None  # Optional collection parameter for SimGrag
    ):
        """
        Initialize the website vectorizer optimized for Kenya Law website.
        
        Args:
            base_url: The root URL of the website to crawl
            model_name: The sentence transformer model to use for embeddings
            chunk_size: Size of text chunks for vectorization
            chunk_overlap: Overlap between chunks
            max_pages: Maximum number of pages to crawl
            max_depth: Maximum depth of links to follow
            vector_db_path: Path to store the vector database
            download_dir: Directory to store downloaded files (if None, uses temp dir)
            concurrent_requests: Maximum number of concurrent requests
            request_delay: Delay between requests (seconds)
            user_agent: Custom user agent string
            respect_robots_txt: Whether to respect robots.txt rules
            proxy: Optional proxy URL
            persist_downloads: Whether to keep downloaded files after processing
            collection: Optional ChromaDB collection (if provided, will use this instead of creating a new one)
        """
        self.base_url = base_url
        self.base_domain = urllib.parse.urlparse(base_url).netloc
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_pages = max_pages
        self.max_depth = max_depth
        self.concurrent_requests = concurrent_requests
        self.request_delay = request_delay
        self.respect_robots_txt = respect_robots_txt
        self.proxy = proxy
        self.persist_downloads = persist_downloads
        
        # Custom user agent rotation for Kenya Law site
        self.user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.110 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/16.4 Safari/605.1.15',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/111.0',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/111.0.0.0 Safari/537.36'
        ]
        self.user_agent = user_agent if user_agent else self.user_agents[0]
        
        # Initialize robots.txt parser
        self.robots_parser = RobotFileParser()
        if self.respect_robots_txt:
            self._init_robots_txt()
        
        # Initialize embedding model
        logger.info(f"Loading embedding model: {model_name}")
        self.embedding_model = SentenceTransformer(model_name)
        
        # Use provided collection or create a new one
        if collection:
            logger.info(f"Using provided collection for {base_url}")
            self.collection = collection
            # Get client from the collection
            self.chroma_client = collection._client
        else:
            # Initialize vector database with optimized settings
            logger.info(f"Initializing vector database at {vector_db_path}")
            self.chroma_client = chromadb.PersistentClient(
                path=vector_db_path,
            )
            
            # Create collection with appropriate settings for search
            collection_name = f"{self.base_domain.replace('.', '_')}_content"
            self.collection = self.chroma_client.get_or_create_collection(
                name=collection_name,
                metadata={"hnsw:space": "cosine"}
            )
        
        # Track visited URLs and downloaded files with better data structures
        self.visited_urls = set()
        self.queued_urls = set()
        self.downloaded_files = set()
        self.failing_urls = set()
        self.stats = CrawlerStats()
        self.stats.start_time = time.time()
        
        # Set up download directory
        if download_dir and persist_downloads:
            self.download_dir = Path(download_dir)
            self.download_dir.mkdir(parents=True, exist_ok=True)
            self.temp_dir = str(self.download_dir)
            self.using_temp_dir = False
        else:
            self.temp_dir = tempfile.mkdtemp()
            self.using_temp_dir = True
            
        # Semaphore for controlling concurrent requests
        self.request_semaphore = asyncio.Semaphore(concurrent_requests)
        
        # Session for better connection pooling
        self.session = None
        
        # Create a file to track crawling progress
        domain_slug = self.base_domain.replace('.', '_')
        self.progress_file = os.path.join(vector_db_path, f"crawl_progress_{domain_slug}.json")
        
        # Kenya Law specific patterns and sections to prioritize
        self.priority_sections = [
            'case_law', 'statutes', 'acts', 'legal_notices', 'bills', 
            'constitution', 'treaties', 'gazette_notices', 'judgments',
            'opinions', 'advisories', 'bench', 'supreme', 'court_of_appeal',
            'high_court', 'law_reports'
        ]
        
        self.pdf_patterns = [
            r'\.pdf$',
            r'download=', 
            r'viewdoc',
            r'view\?',
            r'attachment',
            r'document=',
            r'view_document'
        ]
        
        logger.info(f"WebsiteVectorizer initialized for {base_url}")
        logger.info(f"Download directory: {self.temp_dir} (temporary: {self.using_temp_dir})")
        logger.info(f"Max concurrent requests: {concurrent_requests}, Request delay: {request_delay}s")

    def _init_robots_txt(self):
        """Initialize robots.txt parser for the base domain"""
        try:
            robots_url = urllib.parse.urljoin(self.base_url, "/robots.txt")
            self.robots_parser.set_url(robots_url)
            self.robots_parser.read()
            logger.info(f"Loaded robots.txt from {robots_url}")
        except Exception as e:
            logger.warning(f"Could not read robots.txt: {str(e)}")

    def _can_fetch(self, url: str) -> bool:
        """Check if URL can be fetched according to robots.txt"""
        if not self.respect_robots_txt:
            return True
        
        return self.robots_parser.can_fetch(self.user_agent, url)

    async def _init_session(self):
        """Initialize async HTTP session"""
        import aiohttp
        self.session = aiohttp.ClientSession(
            timeout=aiohttp.ClientTimeout(total=30),
            headers=self._get_headers()
        )

    def _get_headers(self) -> Dict[str, str]:
        """Get request headers with a random user agent"""
        # Rotate user agents to avoid blocking
        user_agent = random.choice(self.user_agents) if self.user_agents else self.user_agent
        
        return {
            'User-Agent': user_agent,
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1',
            'Cache-Control': 'max-age=0',
            'TE': 'Trailers',
        }

    def _save_progress(self):
        """Save crawling progress to file"""
        try:
            progress = {
                "visited_urls": list(self.visited_urls),
                "downloaded_files": list(self.downloaded_files),
                "failing_urls": list(self.failing_urls),
                "stats": asdict(self.stats),
                "last_updated": time.time()
            }
            
            with open(self.progress_file, 'w') as f:
                json.dump(progress, f)
                
            logger.info(f"Saved crawling progress to {self.progress_file}")
        except Exception as e:
            logger.error(f"Error saving progress: {str(e)}")

    def _load_progress(self) -> bool:
        """Load crawling progress from file"""
        try:
            if not os.path.exists(self.progress_file):
                return False
                
            with open(self.progress_file, 'r') as f:
                progress = json.load(f)
                
            self.visited_urls = set(progress.get("visited_urls", []))
            self.downloaded_files = set(progress.get("downloaded_files", []))
            self.failing_urls = set(progress.get("failing_urls", []))
            
            stats = progress.get("stats", {})
            self.stats.pages_visited = stats.get("pages_visited", 0)
            self.stats.files_downloaded = stats.get("files_downloaded", 0)
            self.stats.failed_pages = stats.get("failed_pages", 0)
            self.stats.failed_downloads = stats.get("failed_downloads", 0)
            self.stats.vectorized_texts = stats.get("vectorized_texts", 0)
            
            logger.info(f"Loaded crawling progress from {self.progress_file}")
            logger.info(f"Resuming with {len(self.visited_urls)} visited URLs and "
                       f"{len(self.downloaded_files)} downloaded files")
            return True
            
        except Exception as e:
            logger.error(f"Error loading progress: {str(e)}")
            return False

    async def crawl_and_vectorize(self, resume: bool = True) -> None:
        """
        Main method to crawl website and vectorize all content
        
        Args:
            resume: Whether to resume from previous crawl
        """
        try:
            # Initialize HTTP session
            await self._init_session()
            
            # Try to load previous progress if resume is enabled
            if resume and self._load_progress():
                logger.info(f"Resuming previous crawl")
            else:
                logger.info(f"Starting new crawl of {self.base_url}")
                self.stats.start_time = time.time()
            
            # Store original URL queue in case we need to prioritize sections
            url_queue = asyncio.Queue()
            await url_queue.put((self.base_url, 0))  # (url, depth)
            self.queued_urls.add(self.base_url)
            
            # Use asyncio.gather to limit concurrency with self.request_semaphore
            tasks = []
            for _ in range(self.concurrent_requests):
                tasks.append(asyncio.create_task(self._worker(url_queue)))
                
            # Wait for all workers to complete
            await asyncio.gather(*tasks)
            
            # Close the HTTP session
            if self.session:
                await self.session.close()
                
            elapsed = time.time() - self.stats.start_time
            logger.info(f"Crawl complete. Visited {self.stats.pages_visited} pages and "
                       f"downloaded {self.stats.files_downloaded} files")
            logger.info(f"Failed: {self.stats.failed_pages} pages, {self.stats.failed_downloads} downloads")
            logger.info(f"Vectorized {self.stats.vectorized_texts} text chunks")
            logger.info(f"Total time: {elapsed:.2f} seconds ({elapsed/60:.2f} minutes)")
            
        except Exception as e:
            logger.error(f"Error during crawl and vectorize: {str(e)}")
            raise
        finally:
            # Save final progress
            self._save_progress()
            
            # Close the HTTP session
            if self.session:
                await self.session.close()

    async def _worker(self, queue: asyncio.Queue) -> None:
        """
        Worker process that processes URLs from the queue
        
        Args:
            queue: Queue of URLs to process
        """
        while True:
            try:
                # Get URL from queue with timeout
                try:
                    url, depth = await asyncio.wait_for(queue.get(), timeout=10.0)
                except asyncio.TimeoutError:
                    # If queue is empty for too long, worker exits
                    if queue.empty():
                        break
                    continue
                
                # Process the URL
                await self._process_url(url, depth, queue)
                
                # Mark task as done
                queue.task_done()
                
                # Periodically save progress
                if self.stats.pages_visited % 10 == 0:
                    self._save_progress()
                    self.stats.print_stats()
                    
            except Exception as e:
                logger.error(f"Worker error: {str(e)}")
                
    async def _process_url(self, url: str, depth: int, queue: asyncio.Queue) -> None:
        """
        Process a URL - crawl the page and queue new URLs
        
        Args:
            url: URL to process
            depth: Current depth
            queue: Queue for adding new URLs
        """
        # Skip if we've reached max pages
        if self.stats.pages_visited >= self.max_pages:
            return
            
        # Skip if URL is already visited or is a failing URL
        if url in self.visited_urls or url in self.failing_urls:
            return
            
        # Skip if beyond max depth
        if depth > self.max_depth:
            return
            
        # Skip if not allowed by robots.txt
        if not self._can_fetch(url):
            logger.info(f"Skipping {url} - disallowed by robots.txt")
            return
            
        # Normalize URL for Kenya Law site
        normalized_url = self._normalize_url(url)
        if normalized_url in self.visited_urls:
            return
            
        # Mark as visited
        self.visited_urls.add(normalized_url)
        
        # Add delay between requests
        await asyncio.sleep(self.request_delay)
        
        # Use semaphore to limit concurrent requests
        async with self.request_semaphore:
            try:
                # Log the request
                logger.info(f"Crawling {normalized_url} (depth: {depth})")
                
                # Fetch the page
                page_content, is_file = await self._fetch_url(normalized_url)
                
                if not page_content:
                    logger.warning(f"No content for {normalized_url}")
                    self.stats.failed_pages += 1
                    return
                
                # If it's a downloadable file, process it differently
                if is_file:
                    await self._process_downloadable_file(normalized_url, page_content, content_already_fetched=True)
                    return
                
                # Parse HTML
                soup = BeautifulSoup(page_content, 'html.parser')
                
                # Extract and process main content
                main_content = self._extract_main_content(soup, normalized_url)
                title = soup.title.string if soup.title else "Untitled Page"
                
                # Vectorize the content if significant text was found
                if main_content and len(main_content) > 100:
                    self._vectorize_text(main_content, {
                        "url": normalized_url,
                        "type": "webpage",
                        "title": title.strip() if title else "",
                        "depth": depth
                    })
                
                # Find and queue new URLs to crawl
                await self._extract_and_queue_links(soup, normalized_url, depth, queue)
                
                # Increment counter
                self.stats.pages_visited += 1
                
            except Exception as e:
                logger.error(f"Error processing {normalized_url}: {str(e)}")
                self.failing_urls.add(normalized_url)
                self.stats.failed_pages += 1

    def _normalize_url(self, url: str) -> str:
        """
        Normalize URL to avoid duplicate crawling
        Enhanced for Kenya Law site specifics
        
        Args:
            url: URL to normalize
            
        Returns:
            Normalized URL
        """
        # Parse the URL
        parsed = urllib.parse.urlparse(url)
        
        # Handle Kenya Law specific URL patterns
        if 'kenyalaw.org' in parsed.netloc:
            # Clean up the path - remove trailing slashes, normalize case
            path = parsed.path.rstrip('/')
            
            # Keep important query parameters for PDFs and documents
            if any(pattern in parsed.query.lower() for pattern in ['view', 'download', 'doc', 'file']):
                # Keep the original query string for document files
                normalized = f"{parsed.scheme}://{parsed.netloc}{path}"
                if parsed.query:
                    normalized += f"?{parsed.query}"
                return normalized
                
            # For other pages, ignore the query string
            return f"{parsed.scheme}://{parsed.netloc}{path}"
        
        # Default normalization for other sites
        return f"{parsed.scheme}://{parsed.netloc}{parsed.path.rstrip('/')}"

    async def _fetch_url(self, url: str) -> Tuple[Optional[bytes], bool]:
        """
        Fetch content from a URL using aiohttp
        
        Args:
            url: URL to fetch
            
        Returns:
            Tuple of (content, is_file_download)
        """
        import aiohttp
        
        # Check if URL is downloadable before fetching
        is_file = self._is_downloadable_file(url)
        
        # Retry logic for Kenya Law site - sometimes it rejects initial requests
        max_retries = 3
        retry_delay = 2.0
        retry_count = 0
        
        while retry_count < max_retries:
            try:
                # Get new headers with a random user agent
                headers = self._get_headers()
                
                # Set up proxy if configured
                proxy = self.proxy
                
                # For newer Kenya Law site, skip HEAD request as it may be blocked
                if 'kenyalaw.org' in url:
                    # Directly do GET request
                    async with self.session.get(url, headers=headers, proxy=proxy, 
                                            allow_redirects=True, ssl=False,
                                            timeout=30) as response:
                        
                        # Check status
                        if response.status == 403:
                            logger.warning(f"Access forbidden (403) for {url}, might be rate limited")
                            retry_count += 1
                            await asyncio.sleep(retry_delay * retry_count)
                            continue
                            
                        if response.status != 200:
                            logger.warning(f"Failed to retrieve {url}: HTTP {response.status}")
                            return None, False
                        
                        # Check content type to determine if it's a file
                        content_type = response.headers.get('Content-Type', '')
                        if not is_file and ('pdf' in content_type or 
                                        'octet-stream' in content_type or
                                        'application/' in content_type):
                            is_file = True
                        
                        # Download content
                        content = await response.read()
                        
                        # Double-check if it's actually a PDF file by checking magic bytes
                        if not is_file and len(content) > 4 and content[:4] == b'%PDF':
                            is_file = True
                            logger.info(f"Detected PDF file from header signature: {url}")
                        
                        return content, is_file
                else:
                    # For other sites, first send HEAD request to check content type and size
                    try:
                        async with self.session.head(url, headers=headers, proxy=proxy, 
                                                allow_redirects=True, ssl=False,
                                                timeout=10) as head_response:
                            
                            # Check status
                            if head_response.status != 200:
                                logger.warning(f"HEAD request failed for {url}: {head_response.status}")
                                # Skip HEAD on failure and try direct GET
                                raise aiohttp.ClientError("HEAD request failed, trying direct GET")
                            
                            # Check content type to determine if it's a file
                            content_type = head_response.headers.get('Content-Type', '')
                            if not is_file and ('pdf' in content_type or 
                                            'octet-stream' in content_type or
                                            'application/' in content_type):
                                is_file = True
                            
                            # Check content length to avoid large files
                            content_length = head_response.headers.get('Content-Length')
                            if content_length and int(content_length) > 20 * 1024 * 1024:  # 20 MB limit
                                logger.warning(f"File too large: {url} ({content_length} bytes)")
                                return None, False
                    except (aiohttp.ClientError, asyncio.TimeoutError):
                        # Continue to GET request if HEAD fails
                        pass
                    
                    # Now perform the actual GET request to download content
                    async with self.session.get(url, headers=headers, proxy=proxy, 
                                            allow_redirects=True, ssl=False,
                                            timeout=20) as response:
                        
                        # Check status
                        if response.status != 200:
                            logger.warning(f"Failed to retrieve {url}: HTTP {response.status}")
                            return None, False
                        
                        # Download content
                        content = await response.read()
                        
                        # Double-check if it's actually a PDF file by checking magic bytes
                        if not is_file and len(content) > 4 and content[:4] == b'%PDF':
                            is_file = True
                            logger.info(f"Detected PDF file from header signature: {url}")
                        
                        return content, is_file
                    
            except aiohttp.ClientPayloadError as e:
                logger.error(f"Payload error for {url}: {str(e)}")
                retry_count += 1
                await asyncio.sleep(retry_delay * retry_count)
                
            except aiohttp.ClientError as e:
                logger.error(f"Client error for {url}: {str(e)}")
                return None, False
                
            except asyncio.TimeoutError:
                logger.error(f"Timeout fetching {url}")
                retry_count += 1
                await asyncio.sleep(retry_delay * retry_count)
                
            except Exception as e:
                logger.error(f"Error fetching {url}: {str(e)}")
                return None, False
                
        # If we've exhausted retries
        return None, False

    def _extract_main_content(self, soup: BeautifulSoup, url: str) -> str:
        """
        Extract main content from HTML while removing navigation, headers, footers, etc.
        Optimized for Kenya Law website.
        
        Args:
            soup: BeautifulSoup object
            url: Source URL
            
        Returns:
            Extracted text content
        """
        if 'kenyalaw.org' in url:
            # Kenya Law specific extraction
            main_content = None
            
            # Common content container selectors for Kenya Law site
            content_selectors = [
                'div#main-content', 'div.content-area', 'div.case-content',
                'div.judgment-content', 'article', 'div[role="main"]',
                'main', 'div.page-content', 'div.judgment', 'div.case',
                'div.content', 'div.main', 'div#content', 'div.text-content'
            ]
            
            # Try each selector
            for selector in content_selectors:
                container = soup.select_one(selector)
                if container and len(container.get_text(strip=True)) > 100:
                    main_content = container
                    break
            
            # If no content found with selectors, try common tag-based approach
            if not main_content:
                # Try to find largest text container
                candidates = []
                for elem in soup.find_all(['div', 'article', 'section', 'main']):
                    text_len = len(elem.get_text(strip=True))
                    if text_len > 200:
                        candidates.append((elem, text_len))
                
                if candidates:
                    # Sort by text length and get the largest
                    candidates.sort(key=lambda x: x[1], reverse=True)
                    main_content = candidates[0][0]
            
            if main_content:
                # Remove non-content elements
                for element in main_content.select('nav, header, footer, aside, script, style, noscript, .navigation, .breadcrumbs, .sidebar, .menu, .banner, .header, .footer'):
                    element.decompose()
                
                # Extract text with proper whitespace
                text = main_content.get_text(separator=' ', strip=True)
                lines = [line.strip() for line in text.splitlines() if line.strip()]
                clean_text = '\n'.join(lines)
                
                # Clean up extra whitespace
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                return clean_text.strip()
        
        # Default content extraction for other sites
        # First remove unwanted elements
        for element in soup.select('nav, header, footer, aside, script, style, meta, link, noscript, iframe, form'):
            element.decompose()
        
        # Extract text
        text = soup.get_text(separator=' ', strip=True)
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return '\n'.join(lines)

    def _is_downloadable_file(self, url: str) -> bool:
        """
        Check if URL points to a downloadable file
        Improved detection for Kenya Law documents
        
        Args:
            url: URL to check
            
        Returns:
            True if URL points to downloadable file
        """
        # List of common file extensions
        extensions = ['.pdf', '.doc', '.docx', '.xls', '.xlsx', '.csv', '.txt', '.rtf', '.odt']
        
        # Parse URL
        parsed = urllib.parse.urlparse(url)
        path = parsed.path.lower()
        query = parsed.query.lower()
        
        # Check file extension
        if any(path.endswith(ext) for ext in extensions):
            return True
        
        # Special case for Kenya Law site which uses ASPX with parameters for documents
        if 'kenyalaw.org' in parsed.netloc:
            # Check for PDF patterns in query string
            if any(re.search(pattern, url, re.IGNORECASE) for pattern in self.pdf_patterns):
                return True
            
            # Check for PDF file extension in query string
            if 'pdf' in query or '.pdf' in query:
                return True
            
            # Check for common document view/download parameters
            if any(param in query for param in ['view', 'download', 'file=', 'doc=']):
                return True
        
        return False

    async def _extract_and_queue_links(self, soup: BeautifulSoup, url: str, depth: int, queue: asyncio.Queue) -> None:
        """
        Extract links from page and queue them for processing
        
        Args:
            soup: BeautifulSoup object
            url: Source URL
            depth: Current crawl depth
            queue: Queue for URLs
        """
        # Skip if we've reached max depth
        if depth >= self.max_depth:
            return
            
        links = []
        
        # Find all links
        for link in soup.find_all('a', href=True):
            href = link.get('href', '').strip()
            
            # Skip empty or JavaScript links
            if not href or href.startswith('javascript:') or href == '#':
                continue
                
            # Resolve relative URLs
            link_url = urllib.parse.urljoin(url, href)
            
            # Skip non-HTTP links
            if not link_url.startswith(('http://', 'https://')):
                continue
                
            # Check if same domain
            if not self._is_same_domain(link_url):
                continue
                
            # Normalize the URL
            normalized_url = self._normalize_url(link_url)
            
            # Skip already visited or queued URLs
            if normalized_url in self.visited_urls or normalized_url in self.queued_urls:
                continue
                
            # Mark URL as queued
            self.queued_urls.add(normalized_url)
            
            # Add to links list with priority
            if self._is_downloadable_file(normalized_url):
                priority = 2  # High priority for documents
            elif any(section in normalized_url.lower() for section in self.priority_sections):
                priority = 1  # Medium priority for important sections
            else:
                priority = 0  # Normal priority
                
            links.append((normalized_url, priority))
        
        # Sort links by priority (higher number = higher priority)
        links.sort(key=lambda x: x[1], reverse=True)
        
        # Queue links for processing
        for link_url, _ in links:
            # Add to the queue with increased depth
            await queue.put((link_url, depth + 1))

    def _is_same_domain(self, url: str) -> bool:
        """
        Check if URL is from the same domain as base URL
        
        Args:
            url: URL to check
            
        Returns:
            True if URL is from same domain
        """
        try:
            domain = urllib.parse.urlparse(url).netloc
            return domain == self.base_domain or domain.endswith('.' + self.base_domain) or self.base_domain.endswith('.' + domain)
        except:
            return False

    async def _process_downloadable_file(self, file_url: str, content: bytes = None, content_already_fetched: bool = False) -> None:
        """
        Download and process a file for vectorization
        
        Args:
            file_url: URL of the file to download
            content: Already fetched file content (optional)
            content_already_fetched: Whether content was already fetched
        """
        # Skip if already downloaded
        if file_url in self.downloaded_files:
            return
            
        self.downloaded_files.add(file_url)
        
        try:
            # If content not provided, download it
            if not content_already_fetched or content is None:
                logger.info(f"Downloading file: {file_url}")
                content, _ = await self._fetch_url(file_url)
                
                if not content:
                    logger.warning(f"Failed to download file: {file_url}")
                    self.stats.failed_downloads += 1
                    return
            
            # Determine file type and create file name
            file_ext = self._get_file_extension(file_url, content)
            
            # Generate a filename based on URL hash
            import hashlib
            url_hash = hashlib.md5(file_url.encode()).hexdigest()
            file_name = os.path.join(self.temp_dir, f"{url_hash}{file_ext}")
            
            # Save the file
            with open(file_name, 'wb') as f:
                f.write(content)
            
            # Extract text from file
            text = await self._extract_text_from_file(file_name, file_url)
            
            # Vectorize if text was successfully extracted
            if text and len(text) > 50:  # Ensure we have meaningful content
                self._vectorize_text(text, {
                    "url": file_url, 
                    "type": "file", 
                    "file_type": file_ext[1:],
                    "title": os.path.basename(file_url)
                })
                self.stats.files_downloaded += 1
            else:
                logger.warning(f"No text extracted from {file_url}")
                
            # Clean up temporary file if not persisting downloads
            if not self.persist_downloads:
                os.remove(file_name)
                
        except Exception as e:
            logger.error(f"Error processing file {file_url}: {str(e)}")
            self.stats.failed_downloads += 1

    def _get_file_extension(self, url: str, content: bytes) -> str:
        """
        Determine file extension based on URL and content
        
        Args:
            url: File URL
            content: File content
            
        Returns:
            File extension with dot
        """
        # Check content magic bytes first
        if content[:4] == b'%PDF':
            return '.pdf'
        elif content[:2] == b'PK':
            # Could be docx, xlsx, etc.
            if 'word' in url.lower():
                return '.docx'
            elif 'excel' in url.lower() or 'sheet' in url.lower():
                return '.xlsx'
            else:
                return '.zip'
                
        # Try to get extension from URL
        path = urllib.parse.urlparse(url).path
        ext = os.path.splitext(path)[1].lower()
        
        if ext and ext != '.aspx':
            return ext
            
        # For Kenya Law .aspx files with parameters, check if it's a PDF
        if '.aspx' in url and ('view' in url or 'download' in url):
            return '.pdf'
            
        # Default to .bin if no extension could be determined
        return '.bin'

    async def _extract_text_from_file(self, file_path: str, file_url: str) -> str:
        """
        Extract text from a file based on its type
        
        Args:
            file_path: Path to the downloaded file
            file_url: Original URL of the file
            
        Returns:
            Extracted text content
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            # Handle different file types
            with ThreadPoolExecutor(max_workers=1) as executor:
                if file_extension == '.pdf' or (file_extension == '.aspx' and 'kenyalaw.org' in file_url):
                    # Additional check for PDF disguised as .aspx
                    with open(file_path, 'rb') as f:
                        header = f.read(4)
                        if header.startswith(b'%PDF') or file_extension == '.pdf':
                            # Process as PDF
                            return await asyncio.get_event_loop().run_in_executor(
                                executor, self._extract_text_from_pdf, file_path)
                
                # Handle other file types
                if file_extension in ['.doc', '.docx']:
                    return await asyncio.get_event_loop().run_in_executor(
                        executor, self._extract_text_from_docx, file_path)
                        
                elif file_extension in ['.xls', '.xlsx']:
                    return await asyncio.get_event_loop().run_in_executor(
                        executor, self._extract_text_from_excel, file_path)
                        
                elif file_extension in ['.csv']:
                    return await asyncio.get_event_loop().run_in_executor(
                        executor, self._extract_text_from_csv, file_path)
                        
                elif file_extension in ['.txt', '.md', '.rtf']:
                    return await asyncio.get_event_loop().run_in_executor(
                        executor, self._extract_text_from_text, file_path)
                    
                else:
                    logger.warning(f"Unsupported file type for {file_url}: {file_extension}")
                    return ""
                    
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file with improved error handling
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text = ""
            with open(file_path, 'rb') as f:
                try:
                    pdf_reader = PyPDF2.PdfReader(f)
                    
                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        try:
                            # Try empty password
                            pdf_reader.decrypt('')
                        except:
                            logger.warning(f"Cannot decrypt PDF: {file_path}")
                            return ""
                    
                    # Extract text from each page
                    for page_num in range(len(pdf_reader.pages)):
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n\n"
                        except Exception as e:
                            logger.warning(f"Error extracting text from page {page_num}: {str(e)}")
                    
                except Exception as e:
                    logger.error(f"Error with PyPDF2: {str(e)}")
                    
                    # Fallback to another method if available
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text += page_text + "\n\n"
                    except ImportError:
                        logger.warning("pdfplumber not available for fallback PDF extraction")
                    except Exception as e2:
                        logger.error(f"Error with fallback PDF extraction: {str(e2)}")
            
            # Clean up the text
            text = re.sub(r'\s+', ' ', text).strip()
            
            # If text is too short, it might indicate extraction failure
            if len(text) < 50 and os.path.getsize(file_path) > 5000:
                logger.warning(f"Suspiciously short text extracted from PDF: {file_path}")
                
            return text
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return ""

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
                
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
        return text

    def _extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            # Try with pandas first
            df = pd.read_excel(file_path, sheet_name=None)
            for sheet_name, sheet_df in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_df.to_string(index=False) + "\n\n"
                
        except Exception as e:
            logger.error(f"Error extracting Excel text with pandas: {str(e)}")
            
            # Fallback to openpyxl
            try:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                
                for sheet_name in wb.sheetnames:
                    sheet = wb[sheet_name]
                    text += f"Sheet: {sheet_name}\n"
                    
                    for row in sheet.rows:
                        row_text = " | ".join(str(cell.value) if cell.value is not None else "" for cell in row)
                        text += row_text + "\n"
                    text += "\n"
                    
            except Exception as e2:
                logger.error(f"Error extracting Excel text with openpyxl: {str(e2)}")
                
        return text

    def _extract_text_from_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        text = ""
        try:
            with open(file_path, 'r', newline='', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    text += " | ".join(row) + "\n"
        except Exception as e:
            logger.error(f"Error extracting CSV text: {str(e)}")
            
            # Fallback to pandas
            try:
                df = pd.read_csv(file_path, encoding='utf-8', error_bad_lines=False)
                text = df.to_string(index=False)
            except Exception as e2:
                logger.error(f"Error extracting CSV text with pandas: {str(e2)}")
                
        return text

    def _extract_text_from_text(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
                    
            # If all encodings fail, use binary read and decode with replacement
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='replace')
                
        except Exception as e:
            logger.error(f"Error extracting text: {str(e)}")
            return ""

    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks for vectorization with improved chunking
        
        Args:
            text: Text to chunk
            
        Returns:
            List of text chunks
        """
        if not text or len(text) < 50:
            return []
            
        chunks = []
        
        # Try to split on paragraphs first
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        
        # Process each paragraph
        for para in paragraphs:
            # Skip empty paragraphs
            if not para.strip():
                continue
                
            # If adding this paragraph would exceed chunk size
            if len(current_chunk) + len(para) > self.chunk_size:
                # If current chunk is not empty, add it to chunks
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # If paragraph itself is larger than chunk size, split it
                if len(para) > self.chunk_size:
                    # Split long paragraph into sentences
                    sentences = re.split(r'(?<=[.!?])\s+', para)
                    
                    current_chunk = ""
                    for sentence in sentences:
                        if len(current_chunk) + len(sentence) > self.chunk_size:
                            if current_chunk:
                                chunks.append(current_chunk.strip())
                            
                            # If sentence is too long, force chunk by character
                            if len(sentence) > self.chunk_size:
                                for i in range(0, len(sentence), self.chunk_size - self.chunk_overlap):
                                    chunk = sentence[i:i + self.chunk_size]
                                    if chunk:
                                        chunks.append(chunk.strip())
                            else:
                                current_chunk = sentence
                        else:
                            current_chunk += " " + sentence if current_chunk else sentence
                else:
                    current_chunk = para
            else:
                # Add paragraph with a space
                current_chunk += "\n\n" + para if current_chunk else para
        
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append(current_chunk.strip())
            
        # Create overlap between chunks if they're few and large
        if 1 < len(chunks) <= 3 and all(len(chunk) > self.chunk_size / 2 for chunk in chunks):
            overlapped_chunks = []
            for i in range(len(chunks)):
                if i > 0:
                    # Create an overlapping chunk with end of previous and start of current
                    prev_end = chunks[i-1][-(self.chunk_overlap//2):]
                    curr_start = chunks[i][:self.chunk_size - len(prev_end)]
                    overlapped_chunks.append((prev_end + " " + curr_start).strip())
                    
                overlapped_chunks.append(chunks[i])
                
            chunks = overlapped_chunks
            
        # Filter out empty chunks and deduplicate
        unique_chunks = []
        seen = set()
        
        for chunk in chunks:
            # Normalize for deduplication
            norm_chunk = re.sub(r'\s+', ' ', chunk).strip()
            if norm_chunk and norm_chunk not in seen and len(norm_chunk) > 50:
                seen.add(norm_chunk)
                unique_chunks.append(chunk)
                
        return unique_chunks

    def _vectorize_text(self, text: str, metadata: Dict[str, Any]) -> None:
        """
        Create embeddings for text and store in vector database
        
        Args:
            text: Text to vectorize
            metadata: Metadata for the text
        """
        # Split text into chunks
        chunks = self._chunk_text(text)
        
        if not chunks:
            logger.warning(f"No chunks created for {metadata.get('url', 'unknown URL')}")
            return
            
        logger.info(f"Vectorizing {len(chunks)} chunks from {metadata.get('url', 'unknown')}")
        
        # Create embeddings in batches to avoid memory issues
        batch_size = 16
        
        for i in range(0, len(chunks), batch_size):
            batch_chunks = chunks[i:i+batch_size]
            
            # Create embeddings
            embeddings = self.embedding_model.encode(batch_chunks)
            
            # Create IDs and document metadata
            ids = [f"{metadata['url']}_{i+j}" for j in range(len(batch_chunks))]
            metadatas = [dict(metadata, chunk_index=i+j, chunk_count=len(chunks)) for j in range(len(batch_chunks))]
            
            # Add to vector database
            self.collection.add(
                embeddings=embeddings.tolist(),
                documents=batch_chunks,
                ids=ids,
                metadatas=metadatas
            )
            
        # Update stats
        self.stats.vectorized_texts += len(chunks)

    def query(self, query_text: str, top_k: int = 5, filter_dict: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Query the vector database with filtering options
        
        Args:
            query_text: Text to query
            top_k: Number of results to return
            filter_dict: Optional filter dictionary for metadata filtering
            
        Returns:
            List of results with text and metadata
        """
        # Create embedding for query
        query_embedding = self.embedding_model.encode(query_text).tolist()
        
        # Query the collection
        results = self.collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            include=["documents", "metadatas", "distances"],
            where=filter_dict
        )
        
        # Format results
        formatted_results = []
        for i in range(len(results["documents"][0])):
            formatted_results.append({
                "text": results["documents"][0][i],
                "metadata": results["metadatas"][0][i],
                "distance": results["distances"][0][i]
            })
            
        return formatted_results

    def cleanup(self) -> None:
        """Clean up temporary files"""
        if self.using_temp_dir:
            import shutil
            try:
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temporary directory {self.temp_dir}")
            except Exception as e:
                logger.error(f"Error cleaning up temporary directory: {str(e)}")

class SimGrag:
    """
    Simplified RAG implementation specialized for Kenya Law website
    """
    def __init__(
        self,
        model_name: str = "all-MiniLM-L6-v2",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        vector_db_path: str = "./vector_db",
        context_limit: int = 4000,
        max_context_chunks: int = 10
    ):
        """
        Initialize the SimGrag RAG system for Kenya Law websites
        
        Args:
            model_name: The sentence transformer model to use for embeddings
            chunk_size: Size of text chunks for vectorization
            chunk_overlap: Overlap between chunks
            vector_db_path: Path to store the vector database
            context_limit: Maximum character length for context
            max_context_chunks: Maximum number of context chunks to retrieve
        """
        self.model_name = model_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.vector_db_path = vector_db_path
        self.context_limit = context_limit
        self.max_context_chunks = max_context_chunks
        
        # Initialize embedding model
        logger.info(f"Loading embedding model: {model_name}")
        from sentence_transformers import SentenceTransformer
        self.embedding_model = SentenceTransformer(model_name)
        
        # Initialize vector database client
        logger.info(f"Initializing vector database at {vector_db_path}")
        self.chroma_client = chromadb.PersistentClient(
            path=vector_db_path,
        )
        
        # Create collections for each Kenya Law website
        self.collections = {
            "kenyalaw.org": self.chroma_client.get_or_create_collection(
                name="kenyalaw_org_content",
                metadata={"hnsw:space": "cosine"}
            ),
            "new.kenyalaw.org": self.chroma_client.get_or_create_collection(
                name="new_kenyalaw_org_content",
                metadata={"hnsw:space": "cosine"}
            )
        }
        
        # Initialize vectorizers for each website
        self.vectorizers = {
            "kenyalaw.org": None,
            "new.kenyalaw.org": None
        }
        
    def initialize_vectorizers(self, concurrent_requests: int = 8, request_delay: float = 0.5):
        """
        Initialize website vectorizers for Kenya Law websites
        
        Args:
            concurrent_requests: Maximum number of concurrent requests
            request_delay: Delay between requests (seconds)
        """
        logger.info("Starting vectorizers initialization")
        
        try:
            # Initialize vectorizer for kenyalaw.org
            logger.info("Initializing vectorizer for kenyalaw.org")
            self.vectorizers["kenyalaw.org"] = WebsiteVectorizer(
                base_url="https://kenyalaw.org",
                model_name=self.model_name,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                vector_db_path=self.vector_db_path,
                concurrent_requests=concurrent_requests,
                request_delay=request_delay,
                collection=self.collections["kenyalaw.org"]  # Use the specific collection
            )
            logger.info("Successfully initialized vectorizer for kenyalaw.org")
        except Exception as e:
            logger.error(f"Error initializing kenyalaw.org vectorizer: {str(e)}")
            # Continue execution even if one vectorizer fails
        
        try:
            # Initialize vectorizer for new.kenyalaw.org
            logger.info("Initializing vectorizer for new.kenyalaw.org")
            self.vectorizers["new.kenyalaw.org"] = WebsiteVectorizer(
                base_url="https://new.kenyalaw.org",
                model_name=self.model_name,
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                vector_db_path=self.vector_db_path,
                concurrent_requests=concurrent_requests,
                request_delay=request_delay,
                collection=self.collections["new.kenyalaw.org"]  # Use the specific collection
            )
            logger.info("Successfully initialized vectorizer for new.kenyalaw.org")
        except Exception as e:
            logger.error(f"Error initializing new.kenyalaw.org vectorizer: {str(e)}")
            # Continue execution even if one vectorizer fails
        
        logger.info("Completed vectorizers initialization")
        
    async def crawl_sites(self, max_pages: int = 500, max_depth: int = 5, resume: bool = True):
        """
        Crawl both Kenya Law websites
        
        Args:
            max_pages: Maximum number of pages to crawl per site
            max_depth: Maximum depth of links to follow
            resume: Whether to resume from previous crawl
        """
        for site, vectorizer in self.vectorizers.items():
            if vectorizer is None:
                logger.warning(f"Vectorizer for {site} not initialized. Skipping.")
                continue
                
            # Update max pages and depth
            vectorizer.max_pages = max_pages
            vectorizer.max_depth = max_depth
            
            # Crawl the website
            logger.info(f"Crawling {site}...")
            await vectorizer.crawl_and_vectorize(resume=resume)
            
    def query(self, query_text: str, top_k: int = None, site_filter: str = None) -> List[Dict[str, Any]]:
        """
        Query the vector database across all Kenya Law sites
        
        Args:
            query_text: Text to query
            top_k: Number of results to return per site
            site_filter: Optional site to filter by ("kenyalaw.org" or "new.kenyalaw.org")
            
        Returns:
            List of results with text and metadata
        """
        if top_k is None:
            top_k = self.max_context_chunks
            
        results = []
        sites_to_query = [site_filter] if site_filter else self.collections.keys()
        
        logger.info(f"Querying with text: '{query_text[:50]}...' across sites: {sites_to_query}")
        
        try:
            # Create embedding for query
            query_embedding = self.embedding_model.encode(query_text).tolist()
            
            # Query each collection
            for site in sites_to_query:
                if site not in self.collections:
                    logger.warning(f"Collection for {site} not found. Skipping.")
                    continue
                    
                try:
                    # Query the collection with error handling
                    logger.info(f"Querying collection for {site}")
                    site_results = self.collections[site].query(
                        query_embeddings=[query_embedding],
                        n_results=top_k,
                        include=["documents", "metadatas", "distances"]
                    )
                    
                    # Format and add site results
                    if site_results["documents"] and len(site_results["documents"]) > 0:
                        for i in range(len(site_results["documents"][0])):
                            results.append({
                                "text": site_results["documents"][0][i],
                                "metadata": site_results["metadatas"][0][i],
                                "distance": site_results["distances"][0][i],
                                "site": site
                            })
                        logger.info(f"Found {len(site_results['documents'][0])} results from {site}")
                    else:
                        logger.warning(f"No results found in {site} collection")
                except Exception as e:
                    logger.error(f"Error querying {site} collection: {str(e)}")
                    # Continue with other collections
            
            # Sort combined results by distance (lower is better)
            results.sort(key=lambda x: x["distance"])
            
            # Limit to top_k overall results
            limited_results = results[:top_k]
            logger.info(f"Returning {len(limited_results)} results from query")
            return limited_results
            
        except Exception as e:
            logger.error(f"Error during vector query: {str(e)}")
            # Return empty results on error to prevent API failure
            return []
    
    async def get_response_with_context(self, query: str, top_k: int = None, site_filter: str = None, model_name: str = None):
        """
        Get response with context from both Kenya Law sites
        
        Args:
            query: User query
            top_k: Number of context passages to retrieve (defaults to self.max_context_chunks)
            site_filter: Optional site to filter by ("kenyalaw.org" or "new.kenyalaw.org")
            model_name: Name of the model to use with Ollama (if available)
            
        Returns:
            LLM response with context
        """
        logger.info(f"Getting response for query: '{query[:50]}...'")
        
        try:
            # Use default top_k if not specified
            if top_k is None:
                top_k = self.max_context_chunks
                
            # Query for relevant context across both sites
            context_results = self.query(query, top_k=top_k, site_filter=site_filter)
            
            # Build context string, respecting the context limit
            context_text = ""
            sources = []
            
            for result in context_results:
                # Get source information
                url = result["metadata"].get("url", "unknown")
                title = result["metadata"].get("title", "")
                site = result.get("site", "unknown")
                
                # Track sources for attribution
                if url not in [s[0] for s in sources]:
                    sources.append((url, title))
                
                # Format source with title when available
                source_info = f"Source: {title} ({url})" if title else f"Source: {url}"
                
                # Add text with a separator
                new_context = f"\n\n{source_info}:\n{result['text']}"
                
                # Check if adding this would exceed the context limit
                if len(context_text) + len(new_context) > self.context_limit:
                    # If we're at the limit, stop adding more
                    if context_text:
                        break
                        
                    # If the first context is already too large, truncate it
                    new_context = new_context[:self.context_limit]
                    
                context_text += new_context
            
            # Create prompt for Kenya Law - optimized for LLama3 model
            prompt = f"""<|im_start|>system
You are a Kenya Law Assistant providing accurate information based solely on the Kenya Law website content provided. 
Your role is to assist with queries related to Kenyan laws, statutes, case law, and legal frameworks.

Carefully analyze the following context information to provide accurate answers:

{context_text}

Important guidelines:
1. ONLY use the information provided in the context. Do not rely on prior knowledge.
2. If the context doesn't contain information to answer the question fully, clearly state what information is missing.
3. For legal queries, cite specific sections, cases, or statutes from the context when applicable.
4. Use formal, professional language appropriate for legal discussions.
5. Avoid speculating about legal interpretations beyond what's explicitly stated in the context.
6. When uncertain, acknowledge the limitations of the available information.
7. Keep your response focused and concise.
<|im_end|>

<|im_start|>user
{query}
<|im_end|>

<|im_start|>assistant
"""
            # Try to use Ollama if available
            try:
                import requests
                import os
                import json
                
                # Get Ollama host from environment variable or use default
                ollama_host = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
                logger.info(f"Sending request to Ollama at: {ollama_host}")
                
                # Check if Ollama is available with multiple retries
                max_retries = 3
                retry_delay = 2
                retry_count = 0
                status_response = None
                
                while retry_count < max_retries:
                    try:
                        logger.info(f"Checking Ollama status (attempt {retry_count+1}/{max_retries})...")
                        status_response = requests.get(f"{ollama_host}/api/tags", timeout=5)
                        if status_response.status_code == 200:
                            logger.info("Ollama status check successful")
                            break
                        else:
                            logger.warning(f"Ollama status check failed with code: {status_response.status_code}")
                            retry_count += 1
                            if retry_count < max_retries:
                                logger.info(f"Retrying in {retry_delay} seconds...")
                                await asyncio.sleep(retry_delay)
                                retry_delay *= 2  # Exponential backoff
                    except Exception as e:
                        logger.warning(f"Ollama status check error: {str(e)}")
                        retry_count += 1
                        if retry_count < max_retries:
                            logger.info(f"Retrying in {retry_delay} seconds...")
                            await asyncio.sleep(retry_delay)
                            retry_delay *= 2  # Exponential backoff
                
                if retry_count >= max_retries or not status_response or status_response.status_code != 200:
                    logger.error("Ollama not available after multiple retries")
                    return f"Ollama service is not available. Using context-only response:\n\n{context_text[:500]}...\n\nPlease check that Ollama is running locally or set OLLAMA_HOST environment variable."
                
                # Check if the requested model is available
                try:
                    models = [model.get("name") for model in status_response.json().get("models", [])]
                    if model_name not in models:
                        logger.warning(f"Model {model_name} not found in Ollama. Available models: {models}")
                        # Try to use any available model
                        if models:
                            model_name = models[0]
                            logger.info(f"Falling back to available model: {model_name}")
                        else:
                            return f"Requested model '{model_name}' not available in Ollama. Using context-only response:\n\n{context_text[:500]}..."
                except Exception as status_err:
                    logger.error(f"Error checking Ollama status: {str(status_err)}")
                    # Provide a more helpful error message and fallback response
                    return f"Ollama service not available (Error: {str(status_err)}). Using context-only response:\n\n{context_text[:1000]}..."
                
                # Prepare request payload
                payload = {
                    "model": model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,  # Lower temperature for more factual responses
                        "top_p": 0.95,
                        "top_k": 40
                    ,
                        "num_predict": 1024}
                }
                
                logger.info(f"Sending request to Ollama with model: {model_name}")
                
                # Get timeout from environment variable or use default
                ollama_timeout = int(os.environ.get("OLLAMA_TIMEOUT", 300))
                logger.info(f"Using Ollama timeout of {ollama_timeout} seconds")
                
                # Configure session with timeouts and retry strategy
                session = requests.Session()
                adapter = requests.adapters.HTTPAdapter(
                    max_retries=requests.adapters.Retry(
                        total=3,
                        backoff_factor=1,
                        status_forcelist=[408, 429, 500, 502, 503, 504],
                        allowed_methods=["POST"]
                    )
                )
                session.mount('http://', adapter)
                session.mount('https://', adapter)
                
                # Make request to Ollama API with dynamically set timeout and retry strategy
                response = session.post(
                    f"{ollama_host}/api/generate",
                    json=payload,
                    timeout=(10, ollama_timeout)  # (connect_timeout, read_timeout)
                )
                
                
                if response.status_code != 200:
                    # Check for memory error
                    if response.status_code == 500 and "memory" in response.text.lower():
                        logger.error(f"Memory error with model {model_name}. Returning context-only response.")
                        return f"I found relevant information about your query, but don't have enough memory to generate a full response. Here's the relevant context:\n\n{context_text[:1500]}..."
                    logger.error(f"Error from Ollama API: {response.status_code}, {response.text}")
                    return f"Error accessing Ollama (HTTP {response.status_code}). Using context-only response:\n\n{context_text[:1500]}..."
                if response.status_code == 200:
                    try:
                        result = response.json()
                        logger.info("Successfully received response from Ollama")
                        return result["response"]
                    except json.JSONDecodeError as json_err:
                        logger.error(f"Error decoding Ollama response: {str(json_err)}")
                        return f"Error processing Ollama response. Using context-only response:\n\n{context_text[:500]}..."
                else:
                    logger.error(f"Error from Ollama API: {response.status_code}, {response.text}")
                    return f"Error accessing Ollama (HTTP {response.status_code}). Using context-only response:\n\n{context_text[:500]}...\n\nPlease ensure Ollama is running and the requested model is available."
                    
            except requests.RequestException as req_err:
                logger.error(f"Request error connecting to Ollama: {str(req_err)}")
                return f"Could not connect to Ollama service. Using context-only response:\n\n{context_text[:500]}...\n\nMake sure Ollama is running or set the OLLAMA_HOST environment variable."
            except Exception as e:
                logger.error(f"Unexpected error with Ollama: {str(e)}")
                return f"Error using Ollama LLM: {str(e)}. Using context-only response:\n\n{context_text[:500]}..."
                
        except Exception as e:
            logger.error(f"Error in get_response_with_context: {str(e)}")
            return f"An error occurred while processing your query: {str(e)}"

class LLMContextProvider:
    """
    Class to integrate vectorized content with open source language models
    """
    def __init__(
        self,
        vectorizer: WebsiteVectorizer,
        model_path: str = None,
        context_limit: int = 4000,
        model_type: str = "llama",
        device: str = "cpu",
        max_context_chunks: int = 10
    ):
        """
        Initialize the LLM context provider for open source models
        
        Args:
            vectorizer: WebsiteVectorizer instance
            model_path: Path to the model weights or model ID (HuggingFace)
            context_limit: Maximum character length for context
            model_type: Type of model to use (llama, mistral, falcon, etc.)
            device: Device to run the model on (cpu, cuda, mps)
            max_context_chunks: Maximum number of context chunks to retrieve
        """
        self.vectorizer = vectorizer
        self.model_path = model_path
        self.context_limit = context_limit
        self.model_type = model_type
        self.device = device
        self.max_context_chunks = max_context_chunks
        self.model = None
        self.tokenizer = None
        
        # Load model if path is provided
        if model_path:
            self._load_model()
            
    def _load_model(self):
        """Load the specified open source model"""
        try:
            import torch
            from transformers import AutoModelForCausalLM, AutoTokenizer
            
            logger.info(f"Loading model from {self.model_path} on {self.device}")
            
            # Set appropriate torch dtype based on device
            if self.device == "cuda" and torch.cuda.is_available():
                torch_dtype = torch.float16  # Use half precision for GPU
            else:
                torch_dtype = torch.float32
                
            # Load tokenizer and model
            self.tokenizer = AutoTokenizer.from_pretrained(
                self.model_path,
                use_fast=True
            )
            
            # Load model with appropriate configuration
            self.model = AutoModelForCausalLM.from_pretrained(
                self.model_path,
                torch_dtype=torch_dtype,
                low_cpu_mem_usage=True,
                device_map=self.device
            )
            
            logger.info(f"Model loaded successfully")
            
        except Exception as e:
            logger.error(f"Error loading model: {str(e)}")
            raise
    
    async def get_response_with_context(self, query: str, top_k: int = None, filter_dict: Dict[str, Any] = None) -> str:
        """
        Get response from LLM with relevant context
        
        Args:
            query: User query
            top_k: Number of context passages to retrieve (defaults to self.max_context_chunks)
            filter_dict: Optional filter dictionary for metadata filtering
            
        Returns:
            LLM response with context
        """
        # Use default top_k if not specified
        if top_k is None:
            top_k = self.max_context_chunks
            
        # First, query vector database for relevant context
        context_results = self.vectorizer.query(query, top_k=top_k, filter_dict=filter_dict)
        
        # Build context string, respecting the context limit
        context_text = ""
        sources = []
        
        for result in context_results:
            # Get source information
            url = result["metadata"].get("url", "unknown")
            title = result["metadata"].get("title", "")
            
            # Track sources for attribution
            if url not in [s[0] for s in sources]:
                sources.append((url, title))
            
            # Format source with title when available
            source_info = f"Source: {title} ({url})" if title else f"Source: {url}"
            
            # Add text with a separator
            new_context = f"\n\n{source_info}:\n{result['text']}"
            
            # Check if adding this would exceed the context limit
            if len(context_text) + len(new_context) > self.context_limit:
                # If we're at the limit, stop adding more
                if context_text:
                    break
                    
                # If the first context is already too large, truncate it
                new_context = new_context[:self.context_limit]
                
            context_text += new_context
        
        # Create prompt specifically for Kenya Law - optimized for Llama3
        if 'kenyalaw.org' in self.vectorizer.base_url:
            prompt = f"""<|im_start|>system
You are a Kenya Law Assistant providing accurate information based solely on the Kenya Law website content provided. 
Your role is to assist with queries related to Kenyan laws, statutes, case law, and legal frameworks.

Carefully analyze the following context information to provide accurate answers:

{context_text}

Important guidelines:
1. ONLY use the information provided in the context. Do not rely on prior knowledge.
2. If the context doesn't contain information to answer the question fully, clearly state what information is missing.
3. For legal queries, cite specific sections, cases, or statutes from the context when applicable.
4. Use formal, professional language appropriate for legal discussions.
5. Avoid speculating about legal interpretations beyond what's explicitly stated in the context.
6. When uncertain, acknowledge the limitations of the available information.
7. Keep your response focused and concise.
<|im_end|>

<|im_start|>user
{query}
<|im_end|>

<|im_start|>assistant
"""
        else:
            # Default prompt for other websites - optimized for Llama3
            prompt = f"""<|im_start|>system
You are an AI assistant helping with queries based on the provided content.
Use only the information in the context below to answer the question.

Context information:
{context_text}
<|im_end|>

<|im_start|>user
{query}
<|im_end|>

<|im_start|>assistant
"""
        # If model is loaded, generate a response
        if self.model:
            return await self.generate_response(prompt)
        else:
            # Use external services if available
            if hasattr(self, 'ollama_model'):
                return await self.connect_to_local_ollama(query, top_k=top_k)
            elif hasattr(self, 'api_url'):
                return await self.connect_to_local_api(query, top_k=top_k)
            else:
                # Return just the prompt if no model or service is available
                logger.warning("No model or external service available")
                return "No model loaded. Here are the most relevant passages found:\n\n" + context_text
            
    async def connect_to_local_ollama(self, query: str, model_name: str = "llama3", top_k: int = 5) -> str:
        """
        Connect to a local Ollama instance
        
        Args:
            query: User query
            model_name: Name of the model in Ollama
            top_k: Number of context passages to retrieve
            
        Returns:
            LLM response with context
        """
        try:
            import requests
            
            # Get context
            context_results = self.vectorizer.query(query, top_k=top_k)
            
            # Build better formatted context
            context_parts = []
            for result in context_results:
                url = result["metadata"].get("url", "unknown")
                title = result["metadata"].get("title", "")
                
                # Format source with title when available
                source_info = f"From {title} ({url})" if title else f"From {url}"
                context_parts.append(f"{source_info}:\n{result['text']}")
                
            context_text = "\n\n".join(context_parts)
            
            # Create prompt specifically for Kenya Law
            if 'kenyalaw.org' in self.vectorizer.base_url:
                prompt = f"""
You are a Kenya Law Assistant providing accurate information based solely on the Kenya Law website content provided. 
Your role is to assist with queries related to Kenyan laws, statutes, case law, and legal frameworks.

Carefully analyze the following context information to provide accurate answers:

{context_text}

Important guidelines:
1. ONLY use the information provided in the context. Do not rely on prior knowledge.
2. If the context doesn't contain information to answer the question fully, clearly state what information is missing.
3. For legal queries, cite specific sections, cases, or statutes from the context when applicable.
4. Use formal, professional language appropriate for legal discussions.
5. Avoid speculating about legal interpretations beyond what's explicitly stated in the context.
6. When uncertain, acknowledge the limitations of the available information.

User Question: {query}

Answer:
"""
            else:
                # Default prompt for other websites
                prompt = f"""
Context information:
{context_text}

Given the context information and not prior knowledge, answer the question: {query}
"""
            
            # Make request to Ollama API
            response = requests.post(
                "http://localhost:11434/api/generate",
                json={
                    "model": model_name,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.1,  # Lower temperature for more factual responses
                        "top_p": 0.95,
                        "top_k": 40
                    }
                }
            )
            
            if response.status_code == 200:
                return response.json()["response"]
            else:
                logger.error(f"Error from Ollama API: {response.status_code}, {response.text}")
                return f"Error from Ollama API: {response.status_code}"
                
        except Exception as e:
            logger.error(f"Error connecting to Ollama: {str(e)}")
            return f"Error: {str(e)}"
    
    async def generate_response(self, prompt: str) -> str:
        """Generate response using the loaded model"""
        # Import on demand to avoid unnecessary dependencies
        import torch
        from transformers import pipeline
        
        # Run in executor to avoid blocking
        with ThreadPoolExecutor() as executor:
            future = executor.submit(self._generate_with_transformers_sync, prompt)
            
            # Wait for result
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, future.result)
        
    def _generate_with_transformers_sync(self, prompt: str) -> str:
        """Synchronous generation with transformers"""
        from transformers import pipeline
        
        # Create generation pipeline if not exists
        if not hasattr(self, 'pipe'):
            self.pipe = pipeline(
                "text-generation",
                model=self.model,
                tokenizer=self.tokenizer,
                device=0 if self.device == "cuda" else -1
            )
        
        # Generate
        results = self.pipe(
            prompt,
            max_new_tokens=512,
            temperature=0.7,
            top_p=0.95,
            do_sample=True,
            num_return_sequences=1
        )
        
        # Extract response (remove the prompt)
        generated_text = results[0]['generated_text']
        response = generated_text[len(prompt):]
        
        return response.strip()

async def main():
    import argparse
    
    # Set up command line arguments
    parser = argparse.ArgumentParser(description="Vectorize Kenya Law website and chat with LLM")
    parser.add_argument("--use-simgrag", action="store_true", help="Use SimGrag to source from both kenyalaw.org sites")
    parser.add_argument("--url", type=str, default="https://kenyalaw.org", help="Base URL of the website to crawl (ignored with --use-simgrag)")
    parser.add_argument("--pages", type=int, default=500, help="Maximum number of pages to crawl")
    parser.add_argument("--depth", type=int, default=5, help="Maximum depth of links to follow")
    parser.add_argument("--model", type=str, default=None, help="Path to model or model ID")
    parser.add_argument("--model-type", type=str, default="llama", help="Model type (llama, mistral, falcon, etc.)")
    parser.add_argument("--device", type=str, default="cpu", help="Device to run model on (cpu, cuda, mps)")
    parser.add_argument("--ollama", type=str, default="llama3", help="Use Ollama with specified model name")
    parser.add_argument("--api", type=str, default=None, help="Use local API endpoint")
    parser.add_argument("--no-crawl", action="store_true", help="Skip crawling and use existing vector database")
    parser.add_argument("--query", type=str, default=None, help="Single query mode (no interactive chat)")
    parser.add_argument("--concurrent", type=int, default=8, help="Number of concurrent requests")
    parser.add_argument("--delay", type=float, default=0.5, help="Delay between requests in seconds")
    parser.add_argument("--download-dir", type=str, default=None, help="Directory to store downloaded files")
    parser.add_argument("--persist-downloads", action="store_true", help="Keep downloaded files after processing")
    parser.add_argument("--resume", action="store_true", help="Resume from previous crawl")
    parser.add_argument("--proxy", type=str, default=None, help="HTTP proxy to use")
    parser.add_argument("--site-filter", type=str, default=None, help="Filter SimGrag by site (kenyalaw.org or new.kenyalaw.org)")
    
    args = parser.parse_args()
    
    # Choose between SimGrag and standard WebsiteVectorizer
    if args.use_simgrag:
        # Initialize SimGrag for both Kenya Law sites
        rag = SimGrag(
            vector_db_path="./vector_db",
            chunk_size=1000,
            chunk_overlap=200,
            context_limit=4000,
            max_context_chunks=10
        )
        
        # Initialize vectorizers
        rag.initialize_vectorizers(
            concurrent_requests=args.concurrent,
            request_delay=args.delay
        )
        
        # Perform crawling if not skipped
        if not args.no_crawl:
            logger.info("Crawling both kenyalaw.org and new.kenyalaw.org")
            await rag.crawl_sites(
                max_pages=args.pages,
                max_depth=args.depth,
                resume=args.resume
            )
        else:
            logger.info("Skipping crawl, using existing vector database")
        
        # Single query mode or interactive chat
        if args.query:
            # Process a single query
            logger.info(f"Processing query: {args.query}")
            
            response = await rag.get_response_with_context(
                query=args.query, 
                site_filter=args.site_filter,
                model_name=args.ollama
            )
            
            print("\n" + "="*50)
            print("QUERY:", args.query)
            print("="*50)
            print("RESPONSE:")
            print(response)
            print("="*50 + "\n")
        else:
            # Start interactive chat interface with SimGrag
            chat = SimGragChatInterface(rag, args.site_filter, args.ollama)
            await chat.start()
            
        # No need for cleanup as SimGrag doesn't have a dedicated cleanup method
        
    else:
        # Standard single-site approach
        # Initialize the vectorizer with improved settings
        vectorizer = WebsiteVectorizer(
            base_url=args.url,
            max_pages=args.pages,
            max_depth=args.depth,
            vector_db_path="./vector_db",
            concurrent_requests=args.concurrent,
            request_delay=args.delay,
            download_dir=args.download_dir,
            persist_downloads=args.persist_downloads,
            proxy=args.proxy
        )
        
        # Perform crawling if not skipped
        if not args.no_crawl:
            logger.info(f"Crawling website: {args.url}")
            await vectorizer.crawl_and_vectorize(resume=args.resume)
        else:
            logger.info("Skipping crawl, using existing vector database")
        
        # Initialize LLM context provider
        context_provider = LLMContextProvider(
            vectorizer=vectorizer,
            model_path=args.model,
            model_type=args.model_type,
            device=args.device
        )
        
        # Store API or Ollama model reference for the chat interface
        if args.ollama:
            setattr(context_provider, 'ollama_model', args.ollama)
            logger.info(f"Using Ollama with model: {args.ollama}")
        elif args.api:
            setattr(context_provider, 'api_url', args.api)
            logger.info(f"Using local API at: {args.api}")
        else:
            logger.info(f"Using loaded model: {args.model if args.model else 'None'}")
        
        # Single query mode or interactive chat
        if args.query:
            # Process a single query
            logger.info(f"Processing query: {args.query}")
            
            response = await context_provider.get_response_with_context(args.query)
            
            print("\n" + "="*50)
            print("QUERY:", args.query)
            print("="*50)
            print("RESPONSE:")
            print(response)
            print("="*50 + "\n")
        else:
            # Start interactive chat interface
            chat = ChatInterface(context_provider)
            await chat.start()
        
        # Clean up
        vectorizer.cleanup()

# Interactive chat interface
class ChatInterface:
    """
    Simple interactive chat interface for querying the vectorized content
    """
    def __init__(self, context_provider: LLMContextProvider):
        """Initialize the chat interface"""
        self.context_provider = context_provider
        self.chat_history = []
        
    async def start(self):
        """Start the interactive chat session"""
        print("\n" + "="*50)
        print("Kenya Law Assistant Chat")
        print("="*50)
        print("Type 'exit', 'quit', or 'q' to end the session")
        print("Type 'history' to see previous questions")
        print("Type 'clear' to clear chat history")
        print("="*50 + "\n")
        
        while True:
            try:
                # Get user input
                query = input("\nQuestion: ").strip()
                
                # Check for exit commands
                if query.lower() in ['exit', 'quit', 'q']:
                    print("\nEnding session. Goodbye!")
                    break
                    
                # Check for history command
                if query.lower() == 'history':
                    if not self.chat_history:
                        print("No chat history yet.")
                    else:
                        print("\nChat History:")
                        for i, (q, _) in enumerate(self.chat_history, 1):
                            print(f"{i}. {q}")
                    continue
                    
                # Check for clear command
                if query.lower() == 'clear':
                    self.chat_history = []
                    print("Chat history cleared.")
                    continue
                    
                if not query:
                    continue
                    
                # Process query
                print("Thinking...")
                response = await self.context_provider.get_response_with_context(query)
                
                # Display response
                print("\n" + "-"*50)
                print("Response:")
                print(response)
                print("-"*50)
                
                # Add to history
                self.chat_history.append((query, response))
                
            except KeyboardInterrupt:
                print("\nSession interrupted. Goodbye!")
                break
            except Exception as e:
                print(f"\nError: {str(e)}")

# SimGrag-specific chat interface
class SimGragChatInterface:
    """
    Chat interface for SimGrag RAG system
    """
    def __init__(self, sim_grag: SimGrag, site_filter: str = None, model_name: str = "llama3"):
        """
        Initialize the SimGrag chat interface
        
        Args:
            sim_grag: SimGrag instance
            site_filter: Optional site to filter by ("kenyalaw.org" or "new.kenyalaw.org")
            model_name: Name of the model to use with Ollama
        """
        self.sim_grag = sim_grag
        self.site_filter = site_filter
        self.model_name = model_name
        self.chat_history = []
        
    async def start(self):
        """Start the interactive SimGrag chat session"""
        print("\n" + "="*50)
        print("Kenya Law Assistant Chat (SimGrag)")
        print("="*50)
        if self.site_filter:
            print(f"Currently filtering by site: {self.site_filter}")
        else:
            print("Querying from both kenyalaw.org and new.kenyalaw.org")
            print("Use 'site:kenyalaw.org' or 'site:new.kenyalaw.org' to filter sources")
        print("\nCommands:")
        print("  exit, quit, q - End the session")
        print("  history - See previous questions")
        print("  clear - Clear chat history")
        print("  site:kenyalaw.org - Set site filter to kenyalaw.org")
        print("  site:new.kenyalaw.org - Set site filter to new.kenyalaw.org")
        print("  site:all - Remove site filter")
        print("="*50 + "\n")
        
        while True:
            try:
                # Get user input
                query = input("\nQuestion: ").strip()
                
                # Check for exit commands
                if query.lower() in ['exit', 'quit', 'q']:
                    print("\nEnding session. Goodbye!")
                    break
                    
                # Check for history command
                if query.lower() == 'history':
                    if not self.chat_history:
                        print("No chat history yet.")
                    else:
                        print("\nChat History:")
                        for i, (q, _) in enumerate(self.chat_history, 1):
                            print(f"{i}. {q}")
                    continue
                    
                # Check for clear command
                if query.lower() == 'clear':
                    self.chat_history = []
                    print("Chat history cleared.")
                    continue
                
                # Check for site filter commands
                if query.lower() == 'site:kenyalaw.org':
                    self.site_filter = 'kenyalaw.org'
                    print(f"Set site filter to: {self.site_filter}")
                    continue
                    
                if query.lower() == 'site:new.kenyalaw.org':
                    self.site_filter = 'new.kenyalaw.org'
                    print(f"Set site filter to: {self.site_filter}")
                    continue
                    
                if query.lower() == 'site:all':
                    self.site_filter = None
                    print("Removed site filter. Querying all sites.")
                    continue
                    
                if not query:
                    continue
                
                # Extract site filter from query if present (e.g., "site:kenyalaw.org what is...")
                current_site_filter = self.site_filter
                if query.lower().startswith('site:'):
                    parts = query.split(' ', 1)
                    if len(parts) > 1:
                        site_part = parts[0].lower()
                        if site_part == 'site:kenyalaw.org':
                            current_site_filter = 'kenyalaw.org'
                            query = parts[1]
                        elif site_part == 'site:new.kenyalaw.org':
                            current_site_filter = 'new.kenyalaw.org'
                            query = parts[1]
                        elif site_part == 'site:all':
                            current_site_filter = None
                            query = parts[1]
                    
                # Process query
                print("Thinking...")
                response = await self.sim_grag.get_response_with_context(
                    query=query,
                    site_filter=current_site_filter,
                    model_name=self.model_name
                )
                
                # Display response
                print("\n" + "-"*50)
                print("Response:")
                print(response)
                print("-"*50)
                
                # Add to history
                self.chat_history.append((query, response))
                
            except KeyboardInterrupt:
                print("\nSession interrupted. Goodbye!")
                break
            except Exception as e:
                print(f"\nError: {str(e)}")

if __name__ == "__main__":
    asyncio.run(main())